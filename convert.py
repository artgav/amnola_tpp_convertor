import fitz  # PyMuPDF
import docx
from docx.shared import Pt, RGBColor
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def extract_field(full_text, field_label):
    pattern = rf"{re.escape(field_label)}\s*(.*?)\n"
    match = re.search(pattern, full_text)
    return match.group(1).strip() if match else ""

def add_formatted_text(para, text, size=11, bold=False, italic=False, color=None, underline=False):
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if underline:
        run.font.underline = underline
    if color:
        run.font.color.rgb = color
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    return para


def add_formatted_paragraph(doc, text, size=11, bold=False, italic=False, color=None, underline=False):
    para = doc.add_paragraph()
    return add_formatted_text(para, text, size, bold, italic, color, underline)


def write_kitchen_docx(full_text, output_path):
    doc = docx.Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    event_title = extract_field(full_text, "Event Title:")
    event_date = extract_field(full_text, "Event Worksheet")
    guest_count = extract_field(full_text, "Guest Count:")

    # Header
    add_formatted_paragraph(doc, event_title, bold=True)
    add_formatted_paragraph(doc, f"{event_date}", bold=True)

    sections = []
    cur_section = []
    last_line = ""

    lines = full_text.splitlines()
    for line in lines:
        line = line.strip()
        if line.startswith("Menu Item:") or line.startswith("Beverage Item"):
            if len(cur_section)>0:
               cur_section.pop()
               sections.append(cur_section)
               cur_section = []
            cur_section.append(last_line)
            continue
        last_line = line
        if len(cur_section)<1:
           continue
        #print(line.split('.')[0])
        if line.split('.')[0].isdigit():
           cur_section[-1] += " , " + line
           continue
        if line.startswith("Notes:"):
            continue
        if line.startswith("Quantity:"):
            continue
        if line.startswith("Qty"):
            continue
        if line.startswith("Vendor"):
            continue
        if line.startswith("Miscellaneous"):
            break
        cur_section.append(line)
    
    sections.append(cur_section)

    # Breakfast Header
    for ind,section in enumerate(sections):
     print(section)
     para = doc.add_paragraph()
     add_formatted_text(para,f"{section[0]} ({guest_count} ppl) ",size=16, underline=True)
     if not section[0].startswith("Beverage"):
      
      truck_leaves_time = "TBD"
      try:
        truck_leaves_time = (re.findall(rf"{re.escape('Truck  Leaves')}\s*(.*?)\n", full_text))[ind]
      except:
        pass
      add_formatted_text(para,f"{truck_leaves_time} out",size=16, color=RGBColor(255, 0, 0))

     for line in section[1:]:
        if "features" in line:
            add_formatted_paragraph(doc, line, italic=True)
        else:
            match = re.match(r"(.*?)(\d+\s*ppl)", line)
            if match:
                item, qty = match.groups()
                para = doc.add_paragraph()
                add_formatted_text(para,f"{item.strip()} (")
                add_formatted_text(para,qty.strip(),color=RGBColor(255, 0, 0))
                add_formatted_text(para,f")")
            else:
                add_formatted_paragraph(doc, line, bold=False)

    doc.save(output_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Convert kitchen menu from event PDF to formatted DOCX")
    parser.add_argument("--pdf", required=True, help="Path to input event PDF")
    parser.add_argument("--out", required=True, help="Path to output DOCX file")
    args = parser.parse_args()

    full_text = extract_text_from_pdf(args.pdf)
    write_kitchen_docx(full_text, args.out)
    print(f"Kitchen menu successfully written to {args.out}")
